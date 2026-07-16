from pathlib import Path
from docx import Document


class ChapterExporter:

    def __init__(self):

        self.output = Path("data/exports")

    def export(self, chapter, rows):

        doc = Document()

        doc.add_heading(chapter, level=1)

        for row in rows:

            p = doc.add_paragraph()

            p.add_run(

                f"Q{row['question_no']} "

            ).bold = True

            p.add_run(

                f"({row['marks']} Marks)\n"

            )

            p.add_run(row["question"])

        filename = self.output / f"{chapter}.docx"

        doc.save(filename)