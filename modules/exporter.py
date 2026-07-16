from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
import sqlite3

from modules.database import QUESTION_ORDER_BY


class Exporter:
    """
    Produces the Word question bank. CSV export was dropped: once every
    question is a screenshot, there's no question text to put in a CSV
    anyway, and the Word doc is now the single deliverable.

    Grouping is Subject -> Chapter -> Type (MCQ / Very Short Answer /
    Short Answer / Long Answer / Case Study / Assertion-Reason), using
    the exact same ORDER BY the rest of the app uses (imported from
    database.py, not copy-pasted) so this can never quietly drift out
    of sync with what the app itself shows.
    """

    SELECT_COLUMNS = """
        SELECT
            subject,
            chapter,
            type,
            question_no,
            marks,
            image_path
        FROM questions
    """

    def __init__(self):

        self.db = "data/database/cbse_question_bank.db"

        self.output = Path("exports")
        self.output.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------

    def _fetch_rows(self):

        conn = sqlite3.connect(self.db)

        cursor = conn.cursor()

        cursor.execute(self.SELECT_COLUMNS + QUESTION_ORDER_BY)

        rows = cursor.fetchall()

        conn.close()

        return rows

    # ------------------------------------------------------

    def export(self, df=None):
        """
        Kept as `export()` (not renamed) so app.py doesn't need to
        change its call site. Returns just the Word file path now.
        """

        return self.export_word()

    # ------------------------------------------------------

    def export_word(self, filename="CBSE_Question_Bank.docx"):

        rows = self._fetch_rows()

        doc = Document()

        title = doc.add_heading(
            "CBSE AI Question Bank",
            level=1
        )

        title.runs[0].font.size = Pt(20)

        previous_subject = None
        previous_chapter = None
        previous_type = None

        missing_images = 0

        for row in rows:

            (
                subject,
                chapter,
                q_type,
                question_no,
                marks,
                image_path

            ) = row

            if subject != previous_subject:

                doc.add_page_break()

                doc.add_heading(subject, level=1)

                previous_subject = subject
                previous_chapter = None
                previous_type = None

            if chapter != previous_chapter:

                doc.add_heading(chapter or "Unclassified", level=2)

                previous_chapter = chapter
                previous_type = None

            if q_type != previous_type:

                doc.add_heading(q_type or "Other", level=3)

                previous_type = q_type

            caption = doc.add_paragraph()

            caption.add_run(f"Q{question_no}").bold = True

            if marks:
                caption.add_run(f"  ({marks} Mark{'s' if marks != 1 else ''})")

            if image_path and Path(image_path).exists():

                try:
                    doc.add_picture(str(image_path), width=Inches(5.5))
                except Exception:
                    # Don't let one bad/corrupt image break the whole
                    # export - note it and keep going.
                    doc.add_paragraph("[image could not be inserted]")
                    missing_images += 1

            else:

                doc.add_paragraph("[no screenshot captured for this question]")
                missing_images += 1

        if missing_images:

            doc.add_page_break()
            doc.add_heading("Export Notes", level=1)
            doc.add_paragraph(
                f"{missing_images} question(s) above are missing their "
                f"screenshot image. This usually means the page render "
                f"failed for that specific question - re-processing that "
                f"PDF often resolves it."
            )

        output_file = self.output / filename

        doc.save(output_file)

        return output_file
